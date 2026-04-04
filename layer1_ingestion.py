"""
Layer 1: Ingestion Engine
Parallel document loading, format normalization, secure chunking
"""

import json
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import hashlib

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_EXCEL = True
except:
    HAS_EXCEL = False


@dataclass
class Chunk:
    """Single text chunk with metadata"""
    id: str
    document_id: str
    document_name: str
    content: str
    page_number: int
    section_heading: str
    doc_type: str
    char_offset_start: int
    char_offset_end: int
    chunk_index: int
    total_chunks: int
    ingestion_timestamp: str
    content_hash: str


@dataclass
class Document:
    """Ingested document metadata"""
    id: str
    name: str
    path: str
    doc_type: str
    file_size: int
    page_count: int
    content_preview: str
    ingestion_timestamp: str
    chunk_count: int
    status: str  # INGESTED, CHUNKED, EMBEDDED, FAILED


class IngestionEngine:
    """Document ingestion and chunking"""
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        print("\n[L1] INGESTION ENGINE INITIALIZED")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.documents: Dict[str, Document] = {}
        self.chunks: List[Chunk] = []
    
    def ingest_documents(self, file_paths: List[str]) -> Dict[str, Document]:
        """Load documents from file paths"""
        print(f"\n[L1-INGEST] Loading {len(file_paths)} documents...")
        
        for path_str in file_paths:
            try:
                path = Path(path_str)
                
                if not path.exists():
                    print(f"  ✗ Not found: {path}")
                    continue
                
                if path.is_file():
                    self._ingest_file(path)
                elif path.is_dir():
                    for pattern in ['*.pdf', '*.txt', '*.docx']:
                        for file in path.glob(pattern):
                            if file.is_file():
                                self._ingest_file(file)
            
            except Exception as e:
                print(f"  ✗ Error: {e}")
        
        print(f"[L1-INGEST] ✓ Loaded {len(self.documents)} documents")
        return self.documents
    
    def _ingest_file(self, file_path: Path):
        """Ingest single file"""
        try:
            suffix = file_path.suffix.lower()
            file_size = file_path.stat().st_size
            
            if file_size > 50_000_000:
                print(f"  ⚠ File too large: {file_path.name}")
                return
            
            # Determine document type
            if suffix == '.pdf':
                content, page_count = self._read_pdf(file_path)
                doc_type = 'PDF'
            elif suffix == '.txt':
                content = file_path.read_text(errors='ignore')
                page_count = len(content.split('\n\n')) // 30 + 1
                doc_type = 'TEXT'
            elif suffix == '.docx':
                content, page_count = self._read_docx(file_path)
                doc_type = 'DOCX'
            elif suffix in ['.xlsx', '.xls']:
                content, page_count = self._read_excel(file_path)
                doc_type = 'EXCEL'
            elif suffix == '.csv':
                content = file_path.read_text(errors='ignore')
                page_count = len(content.split('\n'))
                doc_type = 'CSV'
            elif suffix == '.json':
                content = file_path.read_text(errors='ignore')
                page_count = 1
                doc_type = 'JSON'
            else:
                print(f"  ⚠ Unsupported format: {suffix}")
                return
            
            # Create document metadata
            doc_id = hashlib.md5(f"{file_path.name}{datetime.now().isoformat()}".encode()).hexdigest()[:16]
            doc = Document(
                id=doc_id,
                name=file_path.name,
                path=str(file_path),
                doc_type=doc_type,
                file_size=file_size,
                page_count=page_count,
                content_preview=content[:200],
                ingestion_timestamp=datetime.now().isoformat(),
                chunk_count=0,
                status="INGESTED"
            )
            
            self.documents[doc_id] = doc
            print(f"  ✓ {file_path.name} ({doc_type})")
        
        except Exception as e:
            print(f"  ✗ {file_path.name}: {e}")
    
    def _read_pdf(self, path: Path) -> tuple:
        """Extract text from PDF"""
        if not HAS_PDF:
            return "(PDF reading not available)", 1
        
        try:
            reader = PdfReader(path)
            content = ""
            for page in reader.pages:
                content += page.extract_text() + "\n"
            return content[:100_000], len(reader.pages)
        except:
            return "(PDF extraction failed)", 1
    
    def _read_docx(self, path: Path) -> tuple:
        """Extract text from DOCX"""
        if not HAS_DOCX:
            return "(DOCX reading not available)", 1
        
        try:
            doc = DocxDocument(path)
            content = "\n".join([p.text for p in doc.paragraphs])
            return content[:100_000], len(doc.paragraphs) // 20 + 1
        except:
            return "(DOCX extraction failed)", 1
    
    def _read_excel(self, path: Path) -> tuple:
        """Extract text from Excel"""
        if not HAS_EXCEL:
            return "(Excel reading not available)", 1
        
        try:
            dfs = pd.read_excel(path, sheet_name=None)
            content = ""
            for sheet_name, df in list(dfs.items())[:5]:  # First 5 sheets
                content += f"\n[{sheet_name}]\n"
                content += df.to_string()
            return content[:100_000], len(dfs)
        except:
            return "(Excel extraction failed)", 1
    
    def chunk_documents(self):
        """Chunk all ingested documents"""
        print(f"\n[L1-CHUNK] Chunking {len(self.documents)} documents...")
        
        chunk_count = 0
        
        for doc_id, doc in self.documents.items():
            try:
                # Read full content
                full_content = Path(doc.path).read_text(errors='ignore')[:100_000]
                
                # Split into chunks
                chunks = self._create_chunks(full_content, doc_id, doc.name, doc.doc_type)
                self.chunks.extend(chunks)
                
                # Update document
                doc.chunk_count = len(chunks)
                doc.status = "CHUNKED"
                chunk_count += len(chunks)
                
                print(f"  ✓ {doc.name}: {len(chunks)} chunks")
            
            except Exception as e:
                print(f"  ✗ {doc.name}: {e}")
                doc.status = "FAILED"
        
        print(f"[L1-CHUNK] ✓ Created {chunk_count} total chunks")
        return self.chunks
    
    def _create_chunks(self, content: str, doc_id: str, doc_name: str, doc_type: str) -> List[Chunk]:
        """Create overlapping chunks from content"""
        chunks = []
        words = content.split()
        
        chunk_index = 0
        char_pos = 0
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            
            if len(chunk_text.strip()) < 50:  # Skip tiny chunks
                continue
            
            chunk_id = f"{doc_id}_chunk_{chunk_index}"
            content_hash = hashlib.sha256(chunk_text.encode()).hexdigest()
            
            chunk = Chunk(
                id=chunk_id,
                document_id=doc_id,
                document_name=doc_name,
                content=chunk_text,
                page_number=max(1, char_pos // 3000),  # Rough page estimation
                section_heading="",
                doc_type=doc_type,
                char_offset_start=char_pos,
                char_offset_end=char_pos + len(chunk_text),
                chunk_index=chunk_index,
                total_chunks=len(words) // (self.chunk_size - self.chunk_overlap) + 1,
                ingestion_timestamp=datetime.now().isoformat(),
                content_hash=content_hash
            )
            
            chunks.append(chunk)
            chunk_index += 1
            char_pos += len(chunk_text)
        
        return chunks
    
    def save_state(self, output_dir: str = "output"):
        """Save ingestion state to JSON"""
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # Save documents
        docs_data = {
            doc_id: asdict(doc) 
            for doc_id, doc in self.documents.items()
        }
        
        with open(output_path / "documents.json", 'w') as f:
            json.dump(docs_data, f, indent=2)
        
        # Save chunks
        chunks_data = [asdict(chunk) for chunk in self.chunks]
        
        with open(output_path / "chunks.json", 'w') as f:
            json.dump(chunks_data, f, indent=2)
        
        print(f"[L1-SAVE] ✓ Saved to {output_path}")
